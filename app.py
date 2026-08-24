import streamlit as st
import io
import re
import time
import pymupdf
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from google import genai
from google.genai import types

st.set_page_config(page_title="Marathi Document Translator", layout="wide")
st.title("📄 Marathi/Akruti Document ➔ Clean English DOCX")

# Automatic API key detection from Streamlit Secrets
if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    api_key = st.sidebar.text_input("Gemini API Key:", type="password")

def set_cell_background(cell, fill_hex="EAEAEA"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def render_html_table_to_docx(doc, html_table_str):
    clean_html = re.sub(r'<table[^>]*>', '<table>', html_table_str, flags=re.IGNORECASE)
    clean_html = re.sub(r'[\r\n\t]+', ' ', clean_html)
    
    rows = re.findall(r'<tr[^>]*>(.*?)</tr>', clean_html, flags=re.IGNORECASE)
    if not rows:
        return
    
    parsed_grid = []
    for r in rows:
        cells = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', r, flags=re.IGNORECASE)
        if cells:
            parsed_grid.append([re.sub(r'<[^>]+>', '', c).strip() for c in cells])
            
    if not parsed_grid:
        return
        
    num_cols = max(len(r) for r in parsed_grid)
    table = doc.add_table(rows=len(parsed_grid), cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = True
    
    for r_idx, row in enumerate(parsed_grid):
        for c_idx, cell_value in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_value
                if r_idx == 0:
                    set_cell_background(cell, "D9D9D9")
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.name = "Calibri"
                            r.font.size = Pt(10)
                            r.bold = True
                else:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.name = "Calibri"
                            r.font.size = Pt(9.5)

def append_structured_content(doc, content):
    content_no_marathi = re.sub(r'[\u0900-\u097F]+', '', content)
    parts = re.split(r'(<table.*?>.*?</table>)', content_no_marathi, flags=re.DOTALL | re.IGNORECASE)
    
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if part.lower().startswith('<table'):
            render_html_table_to_docx(doc, part)
            doc.add_paragraph()
        else:
            clean_text = re.sub(r'<[^>]+>', '', part)
            for line in clean_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('### '):
                    doc.add_heading(line.replace('### ', ''), level=3)
                elif line.startswith('## '):
                    doc.add_heading(line.replace('## ', ''), level=2)
                elif line.startswith('# '):
                    doc.add_heading(line.replace('# ', ''), level=1)
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(4)

uploaded_file = st.file_uploader("Upload Marathi PDF Document", type=["pdf"])

if uploaded_file and st.button("Translate & Export Clean DOCX"):
    if not api_key:
        st.error("Kripya Gemini API Key set karein.")
    else:
        client = genai.Client(api_key=api_key)
        pdf_doc = pymupdf.open(stream=uploaded_file.read(), filetype="pdf")
        output_doc = Document()
        
        progress = st.progress(0)
        total_pages = len(pdf_doc)
        
        for i in range(total_pages):
            st.write(f"Processing Page {i+1}/{total_pages}...")
            page = pdf_doc[i]
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            
            prompt = """
            You are an expert legal document translator.
            Translate this Marathi document page image into 100% formal ENGLISH.
            
            CRITICAL RULES:
            1. ONLY English Output: Do NOT include original Marathi text, dual language text, or Devanagari script.
            2. Transliterate all proper Marathi names, places, and addresses into English.
            3. IGNORE photos, photo columns, stamps, signatures, and handwritten marginal notes.
            4. FORMAT ALL TABLES USING STANDARD HTML TAGS:
               <table>
                 <tr><th>Sr. No.</th><th>Full Name</th><th>Address</th><th>Designation</th></tr>
                 <tr><td>1</td><td>Name</td><td>Address</td><td>Designation</td></tr>
               </table>
            5. Do not write introduction, notes, or explanations.
            """
            
            max_retries = 3
            translated_text = ""
            for attempt in range(max_retries):
                try:
                    response = client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=[img, prompt]
                    )
                    translated_text = response.text
                    break
                except Exception as e:
                    if attempt < max_retries - 1:
                        time.sleep(3)
                    else:
                        st.error(f"Page {i+1} error: {str(e)}")
            
            if translated_text:
                output_doc.add_heading(f"Page {i+1}", level=1)
                append_structured_content(output_doc, translated_text)
                if i < total_pages - 1:
                    output_doc.add_page_break()
                    
            progress.progress((i + 1) / total_pages)
            time.sleep(1)
            
        out_stream = io.BytesIO()
        output_doc.save(out_stream)
        
        st.success("Translation Complete!")
        st.download_button(
            label="📥 Download Clean English DOCX",
            data=out_stream.getvalue(),
            file_name=f"Clean_English_{uploaded_file.name.replace('.pdf', '')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )